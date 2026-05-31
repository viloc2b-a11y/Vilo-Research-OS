"""
Canonical Native Table Reader for Vilo OS
------------------------------------------
This is the core engine for extracting structural grid data (Schedule of Activities) 
from clinical documents (PDF/DOCX).

Mode Policy:
- 'VALIDATION': Sanitizes all operational identifiers (PHI, Sponsor, Compound) to output safe benchmarking artifacts.
- 'PRODUCTION': Preserves real identifiers (Sponsor, Protocol ID, Site) for true ingestion. 
                Output remains candidate truth and requires human coordinator reconciliation.
"""

import re
import fitz
import docx
import openpyxl
import csv

class NativeTableReader:
    def __init__(self, mode='PRODUCTION'):
        self.mode = mode
        self.phi_keywords = ['zepeda', 'boynton', 'ojeas', 'missy']
        self.sponsor_keywords = ['abbvie', 'abbott', 'acasti', 'allergan', 'gilead', 'paradigm', 'adamis', 'ingenuity', 'coologics', 'boca bio', 'clinica gen bio', 'novartis', 'moderna', 'immunovant']
        self.protocol_keywords = ['para-oa-012', 'para_oa_012', 'mrna-1647', 'imvt-1401', 'crsptl', 'm16-066', 'm14-533', 'udx', 'cgb001', 'app030', 'mv40618', 'rfp_dub-001', 'inception', 'gs-us-553-9020', 'lin-md-64', 'aca-cap-002', 'al 23']
        self.compound_keywords = ['hsv', 'igg', 'lfa', 'ozono', 'nad capilar', 'remdesivir']

    def scrub_text(self, text, doc_id="REDACTED"):
        if self.mode == 'PRODUCTION' or not text:
            return text
            
        # VALIDATION MODE SANITIZATION
        text = str(text)
        text = re.sub(r'[\w\.-]+@[\w\.-]+', '[EMAIL_REDACTED]', text)
        text = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE_REDACTED]', text)
        
        for kw in self.phi_keywords:
            text = re.sub(re.escape(kw), '[STAFF_A]', text, flags=re.IGNORECASE)
        for kw in self.sponsor_keywords:
            text = re.sub(re.escape(kw), '[SPONSOR_A]', text, flags=re.IGNORECASE)
        for kw in self.protocol_keywords:
            text = re.sub(re.escape(kw), f'[{doc_id}]', text, flags=re.IGNORECASE)
        for kw in self.compound_keywords:
            text = re.sub(re.escape(kw), '[INVESTIGATIONAL_PRODUCT_A]', text, flags=re.IGNORECASE)
            
        return text

    def extract_tables(self, file_path, doc_id=""):
        f_lower = file_path.lower()
        if f_lower.endswith('.pdf'):
            return self._extract_pdf(file_path, doc_id)
        elif f_lower.endswith('.docx'):
            return self._extract_docx(file_path, doc_id)
        elif f_lower.endswith('.xlsx'):
            return self._extract_xlsx(file_path, doc_id)
        elif f_lower.endswith('.csv'):
            return self._extract_csv(file_path, doc_id)
        elif f_lower.endswith('.doc') or f_lower.endswith('.xls'):
            raise ValueError(f"Legacy format {file_path[-4:]} unsupported. NEEDS_CONVERSION")
        else:
            raise ValueError("Unsupported file format.")

    def _extract_pdf(self, path, doc_id):
        tables = []
        try:
            doc = fitz.open(path)
            t_idx = 0
            for pnum, page in enumerate(doc):
                tabs = page.find_tables()
                if tabs and tabs.tables:
                    for table in tabs.tables:
                        grid = table.extract()
                        if not grid: continue
                        cells = []
                        for r_idx, row in enumerate(grid):
                            for c_idx, val in enumerate(row):
                                if val is not None:
                                    sanitized_val = self.scrub_text(val.replace('\n', ' '), doc_id)
                                    cells.append({
                                        "row": r_idx,
                                        "col": c_idx,
                                        "rowspan": 1,
                                        "colspan": 1,
                                        "text": sanitized_val,
                                        "bbox": None,
                                        "confidence": None
                                    })
                        tables.append({
                            "table_id": f"{doc_id}_t{t_idx}",
                            "page": pnum + 1,
                            "section": "",
                            "caption": "",
                            "row_count": len(grid),
                            "column_count": len(grid[0]) if len(grid)>0 else 0,
                            "cells": cells,
                            "footnotes": [],
                            "warnings": []
                        })
                        t_idx += 1
            doc.close()
        except Exception as e:
            print(f"Error extracting PDF: {e}")
        return tables

    def _extract_docx(self, path, doc_id):
        tables = []
        try:
            doc = docx.Document(path)
            for t_idx, t in enumerate(doc.tables):
                cells = []
                for r_idx, row in enumerate(t.rows):
                    for c_idx, cell in enumerate(row.cells):
                        sanitized_val = self.scrub_text(cell.text.replace('\n', ' '), doc_id)
                        cells.append({
                            "row": r_idx,
                            "col": c_idx,
                            "rowspan": 1,
                            "colspan": 1,
                            "text": sanitized_val,
                            "bbox": None,
                            "confidence": None
                        })
                tables.append({
                    "table_id": f"{doc_id}_t{t_idx}",
                    "page": 0,
                    "section": "",
                    "caption": "",
                    "row_count": len(t.rows),
                    "column_count": len(t.columns) if t.columns else 0,
                    "cells": cells,
                    "footnotes": [],
                    "warnings": []
                })
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return tables

    def _extract_xlsx(self, path, doc_id):
        tables = []
        try:
            wb = openpyxl.load_workbook(path, data_only=True)
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                cells = []
                max_row = ws.max_row
                max_col = ws.max_column
                for r_idx, row in enumerate(ws.iter_rows()):
                    for c_idx, cell in enumerate(row):
                        if cell.value is not None:
                            val = str(cell.value)
                            sanitized_val = self.scrub_text(val.replace('\n', ' '), doc_id)
                            cells.append({
                                "row": r_idx,
                                "col": c_idx,
                                "rowspan": 1,
                                "colspan": 1,
                                "text": sanitized_val,
                                "bbox": None,
                                "confidence": None
                            })
                tables.append({
                    "table_id": f"{doc_id}_{sheet}",
                    "page": sheet,
                    "section": sheet,
                    "caption": sheet,
                    "row_count": max_row,
                    "column_count": max_col,
                    "cells": cells,
                    "footnotes": [],
                    "warnings": []
                })
        except Exception as e:
            print(f"Error extracting XLSX: {e}")
        return tables

    def _extract_csv(self, path, doc_id):
        tables = []
        try:
            with open(path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                cells = []
                max_col = 0
                r_idx = 0
                for row in reader:
                    max_col = max(max_col, len(row))
                    for c_idx, val in enumerate(row):
                        if val.strip():
                            sanitized_val = self.scrub_text(val.replace('\n', ' '), doc_id)
                            cells.append({
                                "row": r_idx,
                                "col": c_idx,
                                "rowspan": 1,
                                "colspan": 1,
                                "text": sanitized_val,
                                "bbox": None,
                                "confidence": None
                            })
                    r_idx += 1
                tables.append({
                    "table_id": f"{doc_id}_csv",
                    "page": 1,
                    "section": "",
                    "caption": "",
                    "row_count": r_idx,
                    "column_count": max_col,
                    "cells": cells,
                    "footnotes": [],
                    "warnings": []
                })
        except Exception as e:
            print(f"Error extracting CSV: {e}")
        return tables
