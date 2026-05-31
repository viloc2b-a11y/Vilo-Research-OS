import os
import json
import re
import subprocess
import sys

def install(package):
    subprocess.check_call([sys.executable, "-m", "pip", "install", package])

try:
    import fitz # PyMuPDF
except ImportError:
    install('pymupdf')
    import fitz

try:
    import docx
except ImportError:
    install('python-docx')
    import docx

metadata_dir = 'validation-corpus/metadata'
sanitized_dir = 'validation-corpus/sanitized'
registry_file = 'validation-corpus/gold-standard/gold-standard-registry.json'
report_file = os.path.join(metadata_dir, 'safe-text-extraction-batch-1-report.md')

phi_keywords = ['zepeda', 'boynton', 'ojeas', 'missy']
sponsor_keywords = ['abbvie', 'abbott', 'acasti', 'allergan', 'gilead', 'paradigm', 'adamis', 'ingenuity', 'coologics', 'boca bio', 'clinica gen bio']
protocol_keywords = ['para-oa-012', 'para_oa_012', 'mrna-1647', 'imvt-1401', 'crsptl', 'm16-066', 'm14-533', 'udx', 'cgb001', 'app030', 'mv40618', 'rfp_dub-001', 'inception']
compound_keywords = ['hsv', 'igg', 'lfa', 'ozono', 'nad capilar']

def scrub_text(text, doc_id):
    # Regex for emails
    text = re.sub(r'[\w\.-]+@[\w\.-]+', '[EMAIL_REDACTED]', text)
    # Regex for phone-like
    text = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE_REDACTED]', text)
    
    # Generic string replacements
    text_lower = text.lower()
    for kw in phi_keywords:
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub('[STAFF_A]', text)
        
    for kw in sponsor_keywords:
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub('[SPONSOR_A]', text)
        
    for kw in protocol_keywords:
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub(f'[{doc_id}]', text)
        
    for kw in compound_keywords:
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub('[INVESTIGATIONAL_PRODUCT_A]', text)
        
    return text

def check_leaks(text):
    leaks = []
    if '@' in text and '[EMAIL_REDACTED]' not in text: # naive check
        # Look for literal '@' that might be an email, but avoid false positives if possible
        if re.search(r'[\w\.-]+@[\w\.-]+', text):
            leaks.append('Unredacted Email')
    if re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text):
        leaks.append('Unredacted Phone')
        
    text_lower = text.lower()
    for kw in phi_keywords:
        if kw in text_lower: leaks.append(f'PHI: {kw}')
    for kw in sponsor_keywords:
        if kw in text_lower: leaks.append(f'Sponsor: {kw}')
    for kw in protocol_keywords:
        if kw in text_lower: leaks.append(f'Protocol: {kw}')
        
    return leaks

batch1_ids = ['ECRF_GUIDE_A001', 'PROTOCOL_A004_AMEND_001', 'PROTOCOL_A005', 'PROTOCOL_A007', 'PROTOCOL_A009']
mappings = []

for bid in batch1_ids:
    mf = os.path.join(metadata_dir, f"{bid}.mapping.json")
    if os.path.exists(mf):
        with open(mf, 'r', encoding='utf-8') as f:
            mappings.append((mf, json.load(f)))

report_lines = []
report_lines.append("# Safe Text Extraction Batch 1 Report\n")

for mf_path, mapping in mappings:
    if mapping.get('sanitized_artifact_status') != 'PENDING_SAFE_TEXT_EXTRACTION':
        continue
    if mapping.get('usable_for_reader_validation') == True:
        continue
        
    doc_id = mapping['sanitized_id']
    bin_path = mapping.get('unsafe_binary_path', mapping.get('original_relative_path'))
    target_md = mapping['target_path']
    
    report_lines.append(f"### {doc_id}")
    report_lines.append(f"- **Source Binary:** `{bin_path}`")
    
    extracted_text = ""
    pages = 0
    tables = 0
    method = ""
    warnings = []
    
    if bin_path.endswith('.pdf'):
        method = "PyMuPDF (fitz)"
        try:
            doc = fitz.open(bin_path)
            pages = len(doc)
            for page in doc:
                extracted_text += f"\n\n--- Page {page.number + 1} ---\n\n"
                extracted_text += page.get_text("text")
                # PyMuPDF doesn't natively do markdown tables cleanly with just get_text
                # So we count them using find_tables
                tabs = page.find_tables()
                if tabs:
                    tables += len(tabs.tables)
            doc.close()
            
            if len(extracted_text.strip()) < 100 and pages > 0:
                warnings.append("NEEDS_OCR (Image-only PDF detected)")
        except Exception as e:
            warnings.append(f"Extraction Error: {str(e)}")
            
    elif bin_path.endswith('.docx'):
        method = "python-docx"
        try:
            doc = docx.Document(bin_path)
            pages = "N/A"
            tables = len(doc.tables)
            for para in doc.paragraphs:
                extracted_text += para.text + "\n"
            for t in doc.tables:
                extracted_text += "\n[TABLE DETECTED]\n"
                for row in t.rows:
                    extracted_text += " | ".join([c.text.replace('\n', ' ') for c in row.cells]) + "\n"
        except Exception as e:
            warnings.append(f"Extraction Error: {str(e)}")
            
    # Scrub
    scrubbed = scrub_text(extracted_text, doc_id)
    
    # Check Leaks
    leaks = check_leaks(scrubbed)
    
    # Decide status
    status = 'SAFE_TEXT_EXTRACTED'
    usable = True
    
    if leaks or 'NEEDS_OCR' in warnings or 'Extraction Error' in str(warnings):
        status = 'NEEDS_MANUAL_REVIEW'
        usable = False
        warnings.extend(leaks)
        
    # Write output MD
    with open(target_md, 'w', encoding='utf-8') as f:
        f.write(f"# {doc_id}\n\n")
        f.write(f"> **Extraction Method:** {method}\n")
        f.write(f"> **Pages:** {pages}\n")
        f.write(f"> **Tables Detected:** {tables}\n")
        f.write(f"> **Status:** {status}\n\n")
        f.write(scrubbed)
        
    # Update mapping
    mapping['sanitized_artifact_status'] = status
    mapping['usable_for_reader_validation'] = usable
    mapping['extraction_method'] = method
    mapping['extraction_warnings'] = warnings
    mapping['risk_scan_result'] = "CLEAN" if not leaks else "LEAKS_DETECTED"
    
    with open(mf_path, 'w', encoding='utf-8') as f:
        json.dump(mapping, f, indent=2)
        
    # Report
    report_lines.append(f"- **Method:** {method}")
    report_lines.append(f"- **Pages Extracted:** {pages}")
    report_lines.append(f"- **Tables Detected:** {tables}")
    report_lines.append(f"- **Status:** {status}")
    if warnings:
        report_lines.append(f"- **Warnings/Leaks:** {', '.join(warnings)}")
    report_lines.append("")
    
# Update Registry
if os.path.exists(registry_file):
    with open(registry_file, 'r', encoding='utf-8') as f:
        registry = json.load(f)
        
    for doc in registry.get('documents', []):
        if doc['sanitized_id'] in batch1_ids:
            # find the mapping
            mf = os.path.join(metadata_dir, f"{doc['sanitized_id']}.mapping.json")
            if os.path.exists(mf):
                with open(mf, 'r', encoding='utf-8') as fm:
                    mapping = json.load(fm)
                doc['status'] = mapping['sanitized_artifact_status']
                doc['usable_for_reader_validation'] = mapping['usable_for_reader_validation']
            
    with open(registry_file, 'w', encoding='utf-8') as f:
        json.dump(registry, f, indent=2)

# Save report
with open(report_file, 'w', encoding='utf-8') as f:
    f.write('\n'.join(report_lines))

print("Extraction Pipeline completed.")
