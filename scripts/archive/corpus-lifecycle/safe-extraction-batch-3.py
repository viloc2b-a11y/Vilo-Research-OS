import os
import json
import csv
import re
import fitz
import docx

metadata_dir = 'validation-corpus/metadata'
sanitized_dir = 'validation-corpus/sanitized'
registry_file = 'validation-corpus/gold-standard/gold-standard-registry.json'
csv_path = os.path.join(metadata_dir, 'sanitization-plan-batch-3.csv')
report_file = os.path.join(metadata_dir, 'safe-text-extraction-batch-3-report.md')

phi_keywords = ['zepeda', 'boynton', 'ojeas', 'missy']
sponsor_keywords = ['abbvie', 'abbott', 'acasti', 'allergan', 'gilead', 'paradigm', 'adamis', 'ingenuity', 'coologics', 'boca bio', 'clinica gen bio', 'novartis', 'moderna', 'immunovant']
protocol_keywords = ['para-oa-012', 'para_oa_012', 'mrna-1647', 'imvt-1401', 'crsptl', 'm16-066', 'm14-533', 'udx', 'cgb001', 'app030', 'mv40618', 'rfp_dub-001', 'inception', 'gs-us-553-9020', 'lin-md-64', 'aca-cap-002', 'al 23']
compound_keywords = ['hsv', 'igg', 'lfa', 'ozono', 'nad capilar', 'remdesivir']

def scrub_text(text, doc_id):
    text = re.sub(r'[\w\.-]+@[\w\.-]+', '[EMAIL_REDACTED]', text)
    text = re.sub(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', '[PHONE_REDACTED]', text)
    
    for kw in phi_keywords:
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub('[STAFF_A]', text)
        
    for kw in sponsor_keywords:
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub('[SPONSOR_A]', text)
        
    for kw in protocol_keywords:
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub(f'[{doc_id}]', text)
        
    for kw in compound_keywords:
        pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub('[INVESTIGATIONAL_PRODUCT_A]', text)
        
    return text

def check_leaks(text):
    leaks = []
    if '@' in text and '[EMAIL_REDACTED]' not in text:
        if re.search(r'[\w\.-]+@[\w\.-]+', text):
            leaks.append('Unredacted Email')
    if re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text):
        leaks.append('Unredacted Phone')
        
    text_lower = text.lower()
    for kw in phi_keywords:
        if kw.lower() in text_lower: leaks.append(f'PHI: {kw}')
    for kw in sponsor_keywords:
        if kw.lower() in text_lower: leaks.append(f'Sponsor: {kw}')
    for kw in protocol_keywords:
        if kw.lower() in text_lower: leaks.append(f'Protocol: {kw}')
        
    return leaks

targets = []
backups = []
with open(csv_path, 'r', encoding='utf-8') as f:
    reader = csv.DictReader(f)
    for row in reader:
        if row['Status'] == 'BACKUP':
            backups.append(row)
        else:
            targets.append(row)

registry = {}
if os.path.exists(registry_file):
    with open(registry_file, 'r', encoding='utf-8') as f:
        registry = json.load(f)
if 'documents' not in registry:
    registry['documents'] = []

report_lines = []
report_lines.append("# Safe Text Extraction Batch 3 Report\n")

def process_file(c, doc_id_override=None):
    doc_id = doc_id_override if doc_id_override else c['SanitizedID']
    cls = c['DocumentClass']
    
    subfolder = 'protocols'
    if cls == 'AMENDMENT': subfolder = 'amendments'
    elif cls == 'ECRF_GUIDE': subfolder = 'ecrf-guides'
    elif cls == 'PHARMACY_MANUAL': subfolder = 'pharmacy-manuals'
    elif cls == 'LAB_MANUAL': subfolder = 'lab-manuals'
    
    os.makedirs(os.path.join(sanitized_dir, subfolder), exist_ok=True)
    
    bin_path = c['OriginalPath']
    if not os.path.exists(bin_path):
        bin_path = bin_path.replace('\\', '/')
        if not os.path.exists(bin_path):
            bin_path = bin_path.replace('/', os.sep)
            
    target_md = os.path.join(sanitized_dir, subfolder, f"{doc_id}.md")
    
    extracted_text = ""
    pages = 0
    tables = 0
    method = ""
    warnings = []
    status = 'SAFE_TEXT_EXTRACTED'
    usable = True
    
    if not os.path.exists(bin_path):
        warnings.append(f"File not found: {bin_path}")
        method = "Failed"
        status = "NEEDS_MANUAL_REVIEW"
        usable = False
    else:
        if bin_path.endswith('.pdf'):
            method = "PyMuPDF (fitz)"
            try:
                doc = fitz.open(bin_path)
                pages = len(doc)
                for page in doc:
                    extracted_text += f"\n\n--- Page {page.number + 1} ---\n\n"
                    extracted_text += page.get_text("text")
                    tabs = page.find_tables()
                    if tabs: tables += len(tabs.tables)
                doc.close()
                if len(extracted_text.strip()) < 100 and pages > 0:
                    warnings.append("NEEDS_OCR (Image-only PDF detected)")
            except Exception as e:
                warnings.append(f"Extraction Error: {str(e)}")
                
        elif bin_path.endswith('.docx'):
            method = "python-docx"
            try:
                doc = docx.Document(bin_path)
                pages = "N/A"
                tables = len(doc.tables)
                for para in doc.paragraphs:
                    extracted_text += para.text + "\n"
                for t in doc.tables:
                    extracted_text += "\n[TABLE DETECTED]\n"
                    for row in t.rows:
                        extracted_text += " | ".join([ce.text.replace('\n', ' ') for ce in row.cells]) + "\n"
            except Exception as e:
                warnings.append(f"Extraction Error: {str(e)}")
                
        scrubbed = scrub_text(extracted_text, doc_id)
        leaks = check_leaks(scrubbed)
        
        if leaks or 'NEEDS_OCR' in warnings or 'Extraction Error' in str(warnings):
            status = 'NEEDS_MANUAL_REVIEW'
            usable = False
            warnings.extend(leaks)
            
        with open(target_md, 'w', encoding='utf-8') as f:
            f.write(f"# {doc_id}\n\n")
            f.write(f"> **Extraction Method:** {method}\n")
            f.write(f"> **Pages:** {pages}\n")
            f.write(f"> **Tables Detected:** {tables}\n")
            f.write(f"> **Status:** {status}\n\n")
            f.write(scrubbed)
            
    mapping = {
        "sanitized_id": doc_id,
        "original_filename": os.path.basename(bin_path),
        "original_relative_path": bin_path,
        "target_path": target_md,
        "binary_status": "RAW_REFERENCE_ONLY",
        "sanitized_artifact_status": status,
        "usable_for_reader_validation": usable,
        "extraction_method": method,
        "extraction_warnings": warnings,
        "risk_scan_result": "CLEAN" if not warnings else "LEAKS_DETECTED"
    }
    
    mf_path = os.path.join(metadata_dir, f"{doc_id}.mapping.json")
    with open(mf_path, 'w', encoding='utf-8') as f:
        json.dump(mapping, f, indent=2)
        
    registry_entry = {
        "gold_id": doc_id,
        "sanitized_id": doc_id,
        "path": target_md,
        "status": status,
        "usable_for_reader_validation": usable,
        "batch": "BATCH_3"
    }
    
    exists = False
    for i, d in enumerate(registry['documents']):
        if d['sanitized_id'] == doc_id:
            registry['documents'][i] = registry_entry
            exists = True
            break
    if not exists:
        registry['documents'].append(registry_entry)
        
    res = {
        'doc_id': doc_id,
        'bin_path': bin_path,
        'method': method,
        'pages': pages,
        'tables': tables,
        'status': status,
        'warnings': warnings,
        'usable': usable
    }
    return res

final_results = []
used_backups = set()

for c in targets:
    res = process_file(c)
    if not res['usable']:
        # Need backup
        backup_found = False
        for b in backups:
            if b['OriginalPath'] not in used_backups and b['DocumentClass'] == c['DocumentClass']:
                # Found a backup
                res_b = process_file(b, doc_id_override=c['SanitizedID'])
                used_backups.add(b['OriginalPath'])
                backup_found = True
                
                # Report failure of primary
                report_lines.append(f"### {c['SanitizedID']} (FAILED - FAILED OVER TO BACKUP)")
                report_lines.append(f"- **Primary Source:** `{c['OriginalPath']}`")
                report_lines.append(f"- **Failure Reason:** {', '.join(res['warnings'])}")
                
                # Report success of backup
                report_lines.append(f"### {c['SanitizedID']} (BACKUP)")
                report_lines.append(f"- **Source Binary:** `{res_b['bin_path']}`")
                report_lines.append(f"- **Method:** {res_b['method']}")
                report_lines.append(f"- **Pages Extracted:** {res_b['pages']}")
                report_lines.append(f"- **Tables Detected:** {res_b['tables']}")
                report_lines.append(f"- **Status:** {res_b['status']}")
                if res_b['warnings']: report_lines.append(f"- **Warnings/Leaks:** {', '.join(res_b['warnings'])}")
                report_lines.append("")
                break
                
        if not backup_found:
            report_lines.append(f"### {c['SanitizedID']}")
            report_lines.append(f"- **Source Binary:** `{c['OriginalPath']}`")
            report_lines.append(f"- **Status:** {res['status']} (NO BACKUP AVAILABLE)")
            if res['warnings']: report_lines.append(f"- **Warnings/Leaks:** {', '.join(res['warnings'])}")
            report_lines.append("")
    else:
        report_lines.append(f"### {c['SanitizedID']}")
        report_lines.append(f"- **Source Binary:** `{res['bin_path']}`")
        report_lines.append(f"- **Method:** {res['method']}")
        report_lines.append(f"- **Pages Extracted:** {res['pages']}")
        report_lines.append(f"- **Tables Detected:** {res['tables']}")
        report_lines.append(f"- **Status:** {res['status']}")
        report_lines.append("")

with open(registry_file, 'w', encoding='utf-8') as f:
    json.dump(registry, f, indent=2)

with open(report_file, 'w', encoding='utf-8') as f:
    f.write('\n'.join(report_lines))

print("Batch 3 Extraction Pipeline completed.")
